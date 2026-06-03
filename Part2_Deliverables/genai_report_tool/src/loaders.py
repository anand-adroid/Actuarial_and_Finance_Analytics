"""Multi-format document loaders for the RAG corpus.

Handles the real-world mix: markdown (front-matter metadata), PDF (text + tables via
pdfplumber), and Word .docx (paragraphs + tables via python-docx). Metadata for binary
formats comes from a _manifest.csv catalogue (simulating a document-management system).
Production swap: Azure AI Document Intelligence / unstructured.io for parsing at scale.
"""
from __future__ import annotations
import os, glob, csv

def _parse_front_matter(raw: str):
    meta, body = {}, raw
    if raw.startswith("---"):
        end = raw.find("---", 3)
        for line in raw[3:end].strip().splitlines():
            if ":" in line:
                k, v = line.split(":", 1); meta[k.strip()] = v.strip()
        body = raw[end+3:].strip()
    return meta, body

def _load_manifest(corpus_dir: str) -> dict:
    path = os.path.join(corpus_dir, "_manifest.csv")
    out = {}
    if os.path.exists(path):
        for row in csv.DictReader(open(path, newline="")):
            out[row["file"]] = {k: v for k, v in row.items() if k != "file"}
    return out

def _read_pdf(path: str) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
            for tbl in (page.extract_tables() or []):          # flatten tables to text rows
                for row in tbl:
                    parts.append(" | ".join(c or "" for c in row))
    return "\n".join(parts)

def _read_docx(path: str) -> str:
    import docx
    d = docx.Document(path); parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for r in t.rows:
            parts.append(" | ".join(c.text for c in r.cells))
    return "\n".join(parts)

def load_documents(corpus_dir: str):
    """Return list of (doc_id, text, metadata) across .md/.pdf/.docx."""
    manifest = _load_manifest(corpus_dir)
    docs = []
    for path in sorted(glob.glob(os.path.join(corpus_dir, "*"))):
        fn = os.path.basename(path); base, ext = os.path.splitext(fn); ext = ext.lower()
        if fn == "_manifest.csv": continue
        if ext == ".md":
            meta, body = _parse_front_matter(open(path, encoding="utf-8").read())
        elif ext == ".pdf":
            meta, body = dict(manifest.get(fn, {})), _read_pdf(path)
        elif ext == ".docx":
            meta, body = dict(manifest.get(fn, {})), _read_docx(path)
        else:
            continue
        meta["file"] = base; meta["format"] = ext.lstrip(".")
        docs.append((base, body, meta))
    return docs
