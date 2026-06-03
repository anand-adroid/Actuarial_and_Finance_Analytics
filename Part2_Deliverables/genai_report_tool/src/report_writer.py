"""Render the generated report as a branded Word (.docx) document, and optionally PDF.

This is the real-world deliverable: cover page, headers/footers, formatted tables built
from the verified numbers, inline [n] citations, a References section and a sign-off block.
The LLM produced the prose; this layer owns all formatting (so the model never controls layout).
"""
from __future__ import annotations
import os, subprocess, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1F, 0x3B, 0x2E); GREEN = RGBColor(0x2C, 0x5A, 0x3C)

def _style(doc):
    n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)

def render_docx(result, path: str, draft_watermark="DRAFT — pending actuary sign-off"):
    doc = Document(); _style(doc)
    sec = doc.sections[0]; sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
    # footer with page number + draft notice
    footer = sec.footer.paragraphs[0]; footer.text = draft_watermark
    footer.runs[0].font.size = Pt(8); footer.runs[0].font.color.rgb = RGBColor(0x88,0x88,0x88)

    # ---- cover ----
    for _ in range(4): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Actuarial Valuation & Solvency Report"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = NAVY
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f"Long-Term Care Portfolio  |  {result.meta['period']}  |  {result.meta['currency']}")
    rs.font.size = Pt(13); rs.font.color.rgb = GREEN
    note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = note.add_run("Auto-generated draft. Every figure is reconciled to the source of truth; "
                      "pending Appointed Actuary review and sign-off."); rn.italic = True
    rn.font.size = Pt(10); rn.font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_page_break()

    # ---- narrative sections with inline citations ----
    for s in result.sections:
        h = doc.add_heading(s["title"], level=1)
        for run in h.runs: run.font.color.rgb = NAVY
        doc.add_paragraph(s["prose"])

    # ---- tables (from verified numbers) ----
    for tb in result.tables:
        h = doc.add_heading(tb["title"], level=2)
        for run in h.runs: run.font.color.rgb = GREEN
        table = doc.add_table(rows=1, cols=2); table.style = "Light Grid Accent 1"
        hc = table.rows[0].cells; hc[0].text = "Item"; hc[1].text = "Value"
        for label, value in tb["rows"]:
            c = table.add_row().cells; c[0].text = label; c[1].text = str(value)

    # ---- references ----
    h = doc.add_heading("References", level=2)
    for run in h.runs: run.font.color.rgb = GREEN
    for ref in result.references:
        doc.add_paragraph(f"[{ref['n']}] {ref['source']}", style="List Bullet")

    # ---- sign-off ----
    doc.add_heading("Actuary's Statement", level=2)
    doc.add_paragraph("Reviewed and approved by: ____________________________")
    doc.add_paragraph("Appointed Actuary                                   Date: ____________")

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path

def export_pdf(docx_path: str):
    """Best-effort DOCX->PDF. Tries docx2pdf (needs MS Word on Windows/Mac), then LibreOffice."""
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    try:
        from docx2pdf import convert; convert(docx_path, pdf_path); return pdf_path
    except Exception:
        pass
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir",
                        os.path.dirname(docx_path), docx_path], check=False,
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if os.path.exists(pdf_path): return pdf_path
    return None
